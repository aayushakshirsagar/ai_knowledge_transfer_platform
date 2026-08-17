"""File → Markdown → section-tree parsing used by the chunking pipeline.

PDFs are rendered to Markdown with ``pymupdf4llm`` (tables become GFM pipe
tables); DOCX files are converted to Markdown with ``python-docx``.
Standalone images/charts are wrapped into ``<img>`` HTML blocks. The Markdown
is then parsed with ``mistletoe`` (inside an ``HtmlRenderer`` context so
``<img>`` blocks parse as ``HtmlBlock``) into a section tree of our own
``SectionNode`` objects.
"""

from __future__ import annotations

import base64
import re
from dataclasses import dataclass, field
from pathlib import Path

_STANDALONE_IMG_RE = re.compile(r"(?m)^[ \t]*!\[[^\]]*\]\([^)]+\)[ \t]*$")
_IMG_REF_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")

_SUPPORTED_SUFFIXES = {".pdf", ".docx"}


# ---------------------------------------------------------------------------
# Models
# ---------------------------------------------------------------------------


@dataclass
class BlockNode:
    """A leaf content unit inside a section (paragraph, table, chart, ...)."""

    kind: str
    text: str
    isolated: bool = False
    start_line: int | None = None


@dataclass
class SectionNode:
    """A heading and everything that belongs under it."""

    level: int
    title: str
    path: tuple[tuple[int, str], ...] = ()
    blocks: list[BlockNode] = field(default_factory=list)
    children: list[SectionNode] = field(default_factory=list)


@dataclass
class DocTree:
    """Structured document tree built by :func:`build_document_tree`."""

    title: str
    source: str
    root: SectionNode
    source_path: str | None = None


# ---------------------------------------------------------------------------
# File → Markdown
# ---------------------------------------------------------------------------


def _image_ref_to_html(image_ref: str) -> str:
    match = _IMG_REF_RE.match(image_ref.strip())
    if match is None:
        return image_ref
    alt, src = match.groups()
    return f'<img src="{src}" alt="{alt}">'


def _protect_images(markdown: str) -> str:
    """Wrap standalone Markdown image references in ``<img>`` HTML blocks."""
    return _STANDALONE_IMG_RE.sub(lambda m: _image_ref_to_html(m.group(0)), markdown)


def pdf_to_markdown(file_path: str | Path) -> str:
    """Render a PDF to Markdown, embedding images as data URIs."""
    import pymupdf4llm

    return pymupdf4llm.to_markdown(str(file_path), embed_images=True)


def _cell_text(cell) -> str:
    lines: list[str] = []
    for paragraph in cell.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    return "<br>".join(lines) if lines else ""


def _table_to_markdown(table) -> str:
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [_cell_text(cell).replace("|", "\\|") for cell in row.cells]
        rows.append(cells)

    lines: list[str] = []
    header = rows[0] if rows else []
    lines.append("| " + " | ".join(header) + " |")
    lines.append("|" + "|".join(" --- " for _ in header) + "|")
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _paragraph_to_markdown(paragraph, doc) -> str:
    style_name = paragraph.style.name if paragraph.style is not None else "Normal"
    heading_level = _heading_level(style_name)
    text = paragraph.text.strip()

    parts: list[str] = []
    if heading_level:
        parts.append("#" * heading_level + " " + text)
    elif text:
        parts.append(text)

    for blob, content_type in _embedded_images(paragraph, doc):
        data_uri = (
            "data:" + content_type + ";base64," + base64.b64encode(blob).decode("ascii")
        )
        parts.append(f'<img src="{data_uri}" alt="chart">')

    return "\n\n".join(parts)


def _heading_level(style_name: str) -> int | None:
    if style_name == "Title":
        return 1
    if style_name.startswith("Heading"):
        try:
            return max(1, min(int(style_name.rsplit(" ", 1)[-1]), 6))
        except ValueError:
            return 1
    return None


def _embedded_images(paragraph, doc) -> list[tuple[bytes, str]]:
    from docx.oxml.ns import qn

    images: list[tuple[bytes, str]] = []
    for blip in paragraph._p.iter(qn("a:blip")):
        embed_id = blip.get(qn("r:embed"))
        if not embed_id:
            continue
        try:
            part = doc.part.related_parts[embed_id]
        except KeyError:
            continue
        blob = getattr(part, "blob", None)
        if blob:
            images.append((blob, getattr(part, "content_type", "image/png")))
    return images


def docx_to_markdown(file_path: str | Path) -> str:
    """Convert a DOCX file to Markdown using python-docx."""
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = DocxDocument(str(file_path))
    parts: list[str] = []
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            block = _paragraph_to_markdown(Paragraph(child, doc), doc)
        elif child.tag == qn("w:tbl"):
            block = _table_to_markdown(Table(child, doc))
        else:
            continue
        if block:
            parts.append(block)
    return "\n\n".join(parts)


def _to_markdown(file_path: str | Path) -> str:
    suffix = Path(file_path).suffix.lower()
    if suffix == ".pdf":
        markdown = pdf_to_markdown(file_path)
    elif suffix == ".docx":
        markdown = docx_to_markdown(file_path)
    else:
        raise ValueError(
            f"Unsupported file type {suffix!r}; expected one of "
            f"{sorted(_SUPPORTED_SUFFIXES)}"
        )
    return _protect_images(markdown)


# ---------------------------------------------------------------------------
# Markdown → AST → section tree
# ---------------------------------------------------------------------------


def _parse_markdown(markdown: str):
    from mistletoe import Document
    from mistletoe.html_renderer import HtmlRenderer

    with HtmlRenderer():
        return Document(markdown.split("\n"))


def _inline_text(token) -> str:
    """Plain-text representation of an inline (span) token."""
    name = type(token).__name__
    if name in ("RawText", "EscapeSequence"):
        return getattr(token, "content", "")
    if name == "LineBreak":
        return " "
    if name == "InlineCode":
        return getattr(token, "content", "")
    if name == "HtmlSpan":
        return re.sub(r"<[^>]*>", "", getattr(token, "content", ""))
    if name == "AutoLink":
        return getattr(token, "target", "")
    return "".join(_inline_text(c) for c in (getattr(token, "children", None) or []))


def _inline_markdown(token) -> str:
    """Markdown representation of an inline (span) token."""
    name = type(token).__name__
    if name in ("RawText", "EscapeSequence"):
        return getattr(token, "content", "")
    if name == "LineBreak":
        return "\n"
    if name == "InlineCode":
        return "`" + getattr(token, "content", "") + "`"
    if name == "Strong":
        return "**" + "".join(_inline_markdown(c) for c in token.children) + "**"
    if name == "Emphasis":
        return "*" + "".join(_inline_markdown(c) for c in token.children) + "*"
    if name == "Link":
        target = getattr(token, "target", "")
        return "[" + "".join(_inline_markdown(c) for c in token.children) + "](" + target + ")"
    if name == "Image":
        src = getattr(token, "src", "")
        alt = "".join(_inline_markdown(c) for c in (getattr(token, "children", None) or []))
        return f"![{alt}]({src})"
    if name == "AutoLink":
        target = getattr(token, "target", "")
        return f"<{target}>" if target else ""
    if name == "HtmlSpan":
        return getattr(token, "content", "")
    return "".join(_inline_markdown(c) for c in (getattr(token, "children", None) or []))


def _render_table(table) -> str:
    def cell_text(cell) -> str:
        return _inline_text(cell).replace("|", "\\|").strip()

    header = [cell_text(c) for c in table.header.children]
    rows = [[cell_text(c) for c in row.children] for row in table.children]
    ncols = max(len(header), max((len(r) for r in rows), default=0))
    header = header + [""] * (ncols - len(header))
    rows = [row + [""] * (ncols - len(row)) for row in rows]

    lines = ["| " + " | ".join(header) + " |"]
    lines.append("|" + "|".join(" --- " for _ in header) + "|")
    for row in rows:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _render_list(token, depth: int = 0) -> str:
    indent = "    " * depth
    lines: list[str] = []
    ordered = getattr(token, "start", None) is not None
    start = getattr(token, "start", 1) or 1
    for i, item in enumerate(token.children):
        marker = f"{start + i}. " if ordered else "- "
        sub: list[str] = []
        nested: list[str] = []
        for child in getattr(item, "children", []):
            cname = type(child).__name__
            if cname == "Paragraph":
                sub.append(_inline_markdown(child))
            elif cname == "List":
                nested.append(_render_list(child, depth + 1))
            else:
                block = _block_from_token(child)
                if block is not None:
                    sub.append(block.text)
        lines.append(indent + marker + " ".join(sub))
        lines.extend(nested)
    return "\n".join(lines)


def _render_quote(token) -> str:
    parts: list[str] = []
    for child in token.children:
        if type(child).__name__ == "Paragraph":
            parts.append(_inline_markdown(child))
        else:
            block = _block_from_token(child)
            if block is not None:
                parts.append(block.text)
    text = "\n\n".join(parts)
    return "\n".join("> " + line if line else ">" for line in text.split("\n"))


def _block_from_token(token) -> BlockNode | None:
    """Convert a mistletoe block token into a :class:`BlockNode`."""
    name = type(token).__name__
    start = getattr(token, "line_number", None)
    if name == "Paragraph":
        return BlockNode(kind="paragraph", text=_inline_markdown(token), start_line=start)
    if name == "Table":
        return BlockNode(
            kind="table", text=_render_table(token), isolated=True, start_line=start
        )
    if name == "HtmlBlock":
        return BlockNode(
            kind="html_block",
            text=getattr(token, "content", "").strip(),
            isolated=True,
            start_line=start,
        )
    if name in ("CodeFence", "BlockCode"):
        lang = getattr(token, "language", None) or ""
        content = getattr(token, "content", "")
        return BlockNode(kind="code", text=f"```{lang}\n{content}```", start_line=start)
    if name == "List":
        return BlockNode(kind="list", text=_render_list(token), start_line=start)
    if name == "Quote":
        return BlockNode(kind="quote", text=_render_quote(token), start_line=start)
    if name == "ThematicBreak":
        return BlockNode(kind="thematic_break", text="---", start_line=start)
    return None


def _build_section_tree(doc, *, document_title: str | None) -> SectionNode:
    root = SectionNode(level=0, title=document_title or "")
    stack = [root]
    for token in doc.children:
        if type(token).__name__ == "Heading":
            level = getattr(token, "level", 1)
            title = _inline_text(token)
            while stack and stack[-1].level >= level:
                stack.pop()
            parent = stack[-1]
            node = SectionNode(
                level=level, title=title, path=parent.path + ((level, title),)
            )
            parent.children.append(node)
            stack.append(node)
        else:
            block = _block_from_token(token)
            if block is not None:
                stack[-1].blocks.append(block)
    return root


# ---------------------------------------------------------------------------
# Phase 1: document tree
# ---------------------------------------------------------------------------


def build_tree_from_markdown(markdown: str, *, document_title: str | None = None) -> DocTree:
    """Parse Markdown into a :class:`DocTree` (used by tests and direct callers)."""
    protected = _protect_images(markdown)
    document = _parse_markdown(protected)
    root = _build_section_tree(document, document_title=document_title)
    return DocTree(title=document_title or "", source=markdown, root=root)


def build_document_tree(
    file_path: str | Path,
    *,
    document_title: str | None = None,
) -> DocTree:
    """Parse a PDF/DOCX file into a :class:`DocTree`."""
    markdown = _to_markdown(file_path)
    tree = build_tree_from_markdown(markdown, document_title=document_title)
    tree.source_path = str(file_path)
    return tree
