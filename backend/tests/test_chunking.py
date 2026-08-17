from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Inches

from app.ingestion.chunking import chunk_document_tree
from app.ingestion.document_parser import build_document_tree, build_tree_from_markdown

MARKDOWN_WITH_TABLE = """# Quarterly Report

## Revenue

Revenue grew strongly in Q3 across all segments.

| Segment | Q1 | Q2 | Q3 |
| --- | --- | --- | --- |
| North America | 100 | 120 | 150 |
| Europe | 80 | 95 | 110 |
| Asia Pacific | 60 | 70 | 85 |
"""

MARKDOWN_WITH_CHART = """# Expenses

Operating expenses stayed flat.

<img src="data:image/png;base64,AAAA" alt="Expenses chart">
"""


def _tree(markdown: str):
    return build_tree_from_markdown(markdown, document_title="Test")


def test_table_stays_whole_and_isolated() -> None:
    chunks = chunk_document_tree(_tree(MARKDOWN_WITH_TABLE))
    table_chunks = [c for c in chunks if c.chunk_type == "table"]

    assert len(table_chunks) == 1
    body = table_chunks[0].body
    assert body.count("|") == 5 * 5
    assert "North America" in body
    assert "Asia Pacific" in body
    assert table_chunks[0].ancestor_headings == ((1, "Quarterly Report"),)
    assert table_chunks[0].own_heading == (2, "Revenue")


def test_chart_stays_whole_and_isolated() -> None:
    chunks = chunk_document_tree(_tree(MARKDOWN_WITH_CHART))
    chart_chunks = [c for c in chunks if c.chunk_type == "html_block"]

    assert len(chart_chunks) == 1
    assert chart_chunks[0].body.strip().startswith("<img")
    assert chart_chunks[0].own_heading == (1, "Expenses")


def test_build_and_chunk_docx(tmp_path: Path) -> None:
    docx_path = tmp_path / "report.docx"
    doc = DocxDocument()
    doc.add_heading("Quarterly Report", 0)
    doc.add_paragraph("Revenue grew strongly in Q3 across all segments.")
    table = doc.add_table(rows=2, cols=2)
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            cell.text = f"cell-{r}-{c}"
    doc.add_heading("Expenses", level=1)
    doc.add_picture(
        str(Path(__file__).parent / "fixtures" / "chart.png"),
        width=Inches(1),
    )
    doc.save(str(docx_path))

    tree = build_document_tree(docx_path, document_title="Quarterly Report")
    assert tree.title == "Quarterly Report"

    chunks = chunk_document_tree(tree)
    assert any(c.chunk_type == "table" for c in chunks)
    assert any(c.chunk_type == "html_block" for c in chunks)